import requests
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import re

emoji_pattern = re.compile(
    "["
    u"\U0001F600-\U0001F64F"  # emoticons
    u"\U0001F300-\U0001F5FF"  # symbols & pictographs
    u"\U0001F680-\U0001F6FF"  # transport & map symbols
    u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
    u"\U00002702-\U000027B0"
    u"\U000024C2-\U0001F251"
    "]+",
    flags=re.UNICODE,
)


def fetch_from_url(url):
    url_text = requests.get(url)
    soup = BeautifulSoup(url_text.text, 'html.parser')
    response = soup.findAll(['h1', 'p'])
    text = [res.text for res in response]
    texts = ''.join(text)
    texts = texts.replace(".", ".<eos>")
    texts = texts.replace("!", "!<eos>")
    texts = texts.replace("?", "?<eos>")
    texts = texts.replace("\n", "?<eos>")
    texts = texts.replace("<eos>", ' ')
    return texts

def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ''
    for page_num in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page_num].extract_text()
    return text


def read_docx(file_path):
    doc = Document(file_path)

    # Extract text from paragraphs
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # Alternatively, extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    return text


def fetch_from_documents(file):
    file_content = ''
    if file.type == 'text/plain':
        file_content += file.read().decode('utf-8')
    elif file.type == 'application/pdf':
        file_content = read_pdf(file)
    elif file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        file_content = read_docx(file)

    return clean_text(file_content)


def clean_text(x):
    x = x.encode("ascii", "ignore").decode()  # unicode
    x = re.sub(r"https*\S+", " ", x)  # url
    x = re.sub(r"@\S+", " ", x)  # mentions
    x = re.sub(r"#\S+", " ", x)  # hastags
    x = re.sub(r"\s{2,}", " ", x)  # over spaces
    x = emoji_pattern.sub(r"", x)  # emojis
    x = re.sub("[^.,!?A-Za-z0-9]+", " ", x)  # special charachters except .,!?
    return x
